"""
Generate a .docx proposal from the markdown source.
Client-facing content only (stops at the INTERNAL ONLY divider).
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# --- Style configuration ---
style = doc.styles['Normal']
font = style.font
font.name = 'Inter'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Inter'
    heading_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    heading_style.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
    heading_style.paragraph_format.space_after = Pt(6)
    if level == 1:
        heading_style.font.size = Pt(24)
    elif level == 2:
        heading_style.font.size = Pt(18)
    else:
        heading_style.font.size = Pt(14)

# --- Helper functions ---

def add_bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run

def add_paragraph(text, style_name='Normal', bold=False):
    p = doc.add_paragraph(style=style_name)
    # Handle inline bold markers **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def add_blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def add_table(rows, col_widths=None, header=True):
    """Add a formatted table. rows is a list of lists."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

            # Parse bold markers
            parts = re.split(r'(\*\*.*?\*\*)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Inter'
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10)
                    run.font.name = 'Inter'

            # Header row styling
            if header and i == 0:
                for run in p.runs:
                    run.bold = True
                shading = cell._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F2F2F2',
                    qn('w:val'): 'clear'
                })
                shading.append(shading_elem)

    # Light gray borders via table style
    doc.add_paragraph()  # spacing after table
    return table

def add_kv_table(pairs):
    """Add a key-value table (no header row)."""
    table = doc.add_table(rows=len(pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, (key, value) in enumerate(pairs):
        row = table.rows[i]
        # Key cell
        cell_k = row.cells[0]
        cell_k.text = ''
        p = cell_k.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(key)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Inter'

        # Shade key column
        shading = cell_k._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F8F8F8',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elem)

        # Value cell
        cell_v = row.cells[1]
        cell_v.text = ''
        p = cell_v.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.name = 'Inter'

    doc.add_paragraph()
    return table

def add_horizontal_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): 'CCCCCC'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Title
title = doc.add_heading('Progressive Gift Cards', level=1)
subtitle = doc.add_paragraph()
run = subtitle.add_run('Where We Go From Here')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Inter'
subtitle.paragraph_format.space_after = Pt(20)

# Prepared by line
prep = doc.add_paragraph()
run = prep.add_run('Prepared by Redstamp  |  March 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.name = 'Inter'
prep.paragraph_format.space_after = Pt(24)

add_horizontal_rule()

# --- HOW WE GOT HERE ---
doc.add_heading('How We Got Here', level=2)
add_paragraph(
    'Progressive and Redstamp have been working together for some time now \u2014 '
    'from the original operational assessment through to the commerce integration '
    'that replaced phone, email, and fax ordering with a system built around your '
    'margins. In January, we sat down for a wide-ranging conversation about what\u2019s '
    'next. What came out of that call is what this document is built on.'
)

# --- WHERE YOU ARE ---
doc.add_heading('Where You Are', level=2)
add_paragraph('Three things are true about your digital business right now.')

# Subsection 1
doc.add_heading('The fulfillment process can\u2019t scale', level=3)
add_paragraph(
    'Digital card generation runs through one person, one set of scripts, and a manual '
    'pipeline that touches spreadsheets, vendor data, file packaging, and encrypted email '
    'delivery. Loblaws cards alone take 20 seconds each to generate \u2014 a 600-card order '
    'is over three hours.'
)
add_blockquote(
    '\u201cIf you had the same number of digital card orders as you had physical, '
    'then you wouldn\u2019t be able to keep up.\u201d \u2014 Lloyd'
)
add_paragraph(
    'The delivery side compounds this. Encrypted attachments trigger corporate email filters, '
    'customers lose passwords, and every failed delivery creates a support ticket. Fundstream '
    'can get a card to an end user in two seconds. The gap between that and your current process '
    'is widening.'
)
add_paragraph(
    'The back-office friction is just as real. Every order that comes through Formidable Forms '
    'generates an email notification \u2014 and from there, someone has to manually create the '
    'corresponding invoice in QuickBooks. There\u2019s no sync, no automation, just a person '
    're-entering data from one system into another. It works at current volume, but it\u2019s the '
    'kind of process that quietly breaks as the business grows.'
)

# Subsection 2
doc.add_heading('You\u2019re committed to change \u2014 and you need the right partner to navigate it', level=3)
add_paragraph(
    'You\u2019ve already started moving: the sticker machine for physical cards, barcode capture '
    'for inventory, merchant backend integrations on a 3-5 year horizon. You\u2019ve standardized '
    'on URL-based delivery to simplify automation. Customer self-service for digital card retrieval '
    'is a top priority. Branded vendor experiences for merchants like Save-on-Foods and Sequoia '
    'represent an opportunity to expand \u2014 taking over fulfillment for stores and chains that '
    'are currently handling it themselves.'
)
add_paragraph(
    'You\u2019re committed to the work and effort these organizational and operational changes require. '
    'We\u2019ve been in this together since before you had an online ordering system \u2014 when every '
    'order came in by phone, email, or fax. The commerce integration we built has processed hundreds '
    'of thousands of dollars in orders through a model that fits your margins, not a generic checkout '
    'flow. That foundation got you here. Now the business has reached the point where you\u2019re ready '
    'to see what the next phase looks like and how to move forward. Your operation is one of two of '
    'its kind in Canada \u2014 there\u2019s no off-the-shelf playbook. Each step needs to build on '
    'the last so that progress compounds instead of competing for attention.'
)
add_blockquote(
    '\u201cI want you to come back and tell me what we should be doing. From your basis of your '
    'expertise. Because I\u2019m happy to invest in you guys to help us make this better.\u201d \u2014 Doug'
)

# Subsection 3
doc.add_heading('The team is ready for this to change', level=3)
add_paragraph(
    'Lloyd has been direct about wanting his role automated. Your team understands the risk of '
    'having critical processes concentrated in one person\u2019s tooling. The willingness to move '
    'is real \u2014 what you need is someone to translate that willingness into architecture, '
    'timelines, and a plan your team can execute against.'
)
add_paragraph('That\u2019s what this document lays out.')

add_horizontal_rule()

# --- WHAT WE RECOMMEND ---
doc.add_heading('What We Recommend', level=2)
p = add_paragraph('**Solve the card generation problem first. Then rebuild the order platform on top of it.**')
add_paragraph(
    'The instinct is to start with the customer experience \u2014 it\u2019s the most visible pain '
    'point and the thing your customers would notice immediately. But a new front door without a '
    'reliable fulfillment engine behind it is a storefront with nothing on the shelves. The generation '
    'process is the bottleneck that constrains everything else: delivery speed, staff flexibility, '
    'volume capacity, and your ability to compete with Fundstream on digital.'
)
add_paragraph(
    'Digital gift card volume is growing across the industry. Your direct vendor relationships \u2014 '
    'one of only two Walmart bulk distributors in Canada \u2014 are the competitive advantage no one '
    'can replicate. But that advantage only holds if your fulfillment infrastructure can keep pace '
    'with demand. Right now, it can\u2019t.'
)
add_paragraph('We recommend a phased approach where each step builds on the last and delivers standalone value.')

# Phase 1
doc.add_heading('Phase 1 \u2014 Capture what exists', level=3)
add_paragraph(
    'Before we build anything, we need to fully understand what Lloyd has built and how it works '
    'across every vendor. This is not a discovery exercise \u2014 we know what needs to happen. '
    'It\u2019s a structured hand-off session.'
)
add_paragraph(
    'Danny, our Director of Operations, leads a half-day working session with Lloyd at your office '
    '\u2014 approximately four hours, with Spencer joining remotely and a member of the technical '
    'team available as needed. The goal is complete documentation of every vendor workflow: what data '
    'comes in, what scripts run, what the output looks like, where things break, and what the edge '
    'cases are. Walmart, Amazon, Chapters, Loblaws \u2014 each one mapped step by step.'
)
add_paragraph('This session produces two things:')
p = doc.add_paragraph(style='List Number')
add_bold_text(p, 'Vendor runbooks')
p.add_run(
    ' \u2014 step-by-step documentation written for operators, not developers. Detailed enough '
    'that Mario or a new hire could reference them today, even before the new tool is built.'
)
p = doc.add_paragraph(style='List Number')
add_bold_text(p, 'A technical specification')
p.add_run(
    ' for the card generation tool \u2014 informed by what actually happens in practice, not assumptions.'
)
add_paragraph(
    'We can also revisit the order matching work during this phase. The tooling available today '
    'makes that problem significantly more straightforward than when we first looked at it.'
)
add_kv_table([
    ('Timeline', 'April 2026'),
    ('Investment', '[To be confirmed with team]'),
])

# Phase 2
doc.add_heading('Phase 2 \u2014 Build the card generation tool', level=3)
add_paragraph(
    'This is the core build. A secure application that replaces Lloyd\u2019s scripts with a system '
    'any trained staff member can operate.'
)
add_paragraph(
    'The goal is simple: Mario selects a vendor, uploads the card data, and the system handles '
    'generation, validation, and URL output. The same process Lloyd runs today \u2014 without the '
    'spreadsheets, manual scripting, and single-point-of-failure risk.'
)
add_table([
    ['What changes', 'What it means'],
    ['Any trained staff member can process orders', 'No more single-point-of-failure dependency on Lloyd'],
    ['New hire training in under a day', 'The system guides the workflow, not tribal knowledge'],
    ['New vendors added through configuration', 'No custom scripting project for each merchant'],
    ['Secure handling from intake through delivery', 'Card data protected throughout the process'],
])
add_paragraph(
    'We\u2019re confident in delivering a solution that solves this long-term. One that\u2019s built '
    'to support future growth, additional vendors, and new features as your business evolves. The '
    'mass distribution capability \u2014 sending cards directly to hundreds of individual recipients '
    '\u2014 is a natural extension that can come later. The priority is eliminating the single point '
    'of failure and building a foundation that holds.'
)
add_paragraph(
    'A traditional software firm starting from scratch would scope this type of custom application at '
    '$55K-$120K. We\u2019re not starting from scratch. Redstamp already understands your business, '
    'your vendors, and your operational constraints \u2014 context that would take an outside firm '
    'weeks of billable discovery to acquire. Combined with the development infrastructure we\u2019ve '
    'built internally to work efficiently at our scale, we can deliver this at a fraction of the '
    'market rate without cutting corners on security or quality.'
)
add_kv_table([
    ('Timeline', 'May\u2013September 2026'),
    ('Market rate', '$55K\u2013$120K'),
    ('Redstamp investment', '[To be confirmed with team]'),
])
add_paragraph(
    '**Phase 2 is designed to ship before the holiday busy season** \u2014 the card generation tool '
    'needs to be live and stable before order volume ramps up in Q4, which means complete and in '
    'your team\u2019s hands no later than early October.'
)

# Phase 3
doc.add_heading('Phase 3 \u2014 Rebuild the order platform', level=3)
add_paragraph(
    'The Formidable Forms setup has served its purpose. It proved the model works, and hundreds of '
    'thousands of dollars in orders have moved through it. But you\u2019re processing six-figure gift '
    'card orders through a form plugin with no customer accounts, no order history, and no direct '
    'connection to your back-office systems. As order volume grows and the business matures, the '
    'platform needs to reflect that.'
)
add_paragraph(
    'Your marketing site stays exactly where it is \u2014 WordPress continues to serve that purpose '
    'well, and there\u2019s no reason to disrupt it. The order platform is a separate application, '
    'purpose-built for commerce, that lives alongside the marketing site rather than being bolted onto it.'
)
add_table([
    ['Capability', 'What it replaces'],
    ['Customer accounts with 2FA', 'No accounts \u2014 orders placed through an anonymous form'],
    ['Streamlined order submission', 'Formidable Forms plugin handling six-figure transactions'],
    ['Order history and status tracking', 'Email threads and manual follow-ups'],
    ['Admin backend', 'Scattered spreadsheets and email notifications'],
    ['Direct QuickBooks integration', 'Manual invoice creation from form submission emails'],
])
add_paragraph(
    'Benji Pays stays in place for payment capture \u2014 the platform handles everything around '
    'the payment without taking on credit card processing or the margin compression that comes with it.'
)
add_paragraph('This phase is the largest investment, and we\u2019d approach it in stages:')

# Stages as mini sections
p = doc.add_paragraph()
add_bold_text(p, 'Stage 1 \u2014 MVP')
p.add_run(
    '  |  Customer accounts, 2FA, order submission, status tracking, admin view, and QuickBooks '
    'integration. This replaces Formidable Forms, eliminates the manual invoice creation bottleneck, '
    'and establishes the new foundation.'
)

p = doc.add_paragraph()
add_bold_text(p, 'Stage 2 \u2014 Digital delivery')
p.add_run(
    '  |  Digital card retrieval connected to the generation tool, direct-to-recipient distribution. '
    'This is where the card generation tool and the order platform come together \u2014 customers '
    'retrieve cards through the platform instead of encrypted email.'
)

p = doc.add_paragraph()
add_bold_text(p, 'Stage 3 \u2014 White-label')
p.add_run(
    '  |  Branded vendor portals. Standing up a Save-on-Foods or Sequoia instance becomes a '
    'configuration exercise, not a rebuild. We\u2019ll architect for this from day one, but won\u2019t '
    'build it until the contracts justify it.'
)

doc.add_paragraph()  # spacing

add_paragraph(
    'A traditional firm would scope a custom B2B platform of this complexity at $90K-$210K. The '
    'same advantages apply here \u2014 our existing knowledge of your operations, the shared '
    'infrastructure with the card generation tool, and efficient development workflows mean this '
    'investment will be significantly below market rate.'
)

add_table([
    ['Stage', 'Timeline', 'Market rate', 'Redstamp investment'],
    ['**Stage 1** \u2014 MVP', 'Q4 2026\u2013Q1 2027', '$90K\u2013$210K (full platform)', '[To be confirmed with team]'],
    ['**Stage 2** \u2014 Digital delivery', 'Q1\u2013Q2 2027', '(included above)', '[Scoped once MVP is live]'],
    ['**Stage 3** \u2014 White-label', '2027, when contracts justify it', '(included above)', '[Scoped when contracts justify it]'],
])

add_horizontal_rule()

# --- WHAT THIS REQUIRES ---
doc.add_heading('What This Requires', level=2)
add_paragraph(
    'We want to be straightforward about what this engagement looks like \u2014 both the build '
    'and what comes after.'
)

doc.add_heading('The build', level=3)
add_paragraph(
    'Each phase represents a discrete investment. You don\u2019t commit to everything upfront \u2014 '
    'you see the results of one phase before deciding on the next. Timelines assume an early April '
    'start and will sharpen as each phase completes. We\u2019ll firm up Phase 2 pricing after the '
    'Phase 1 session gives us full visibility into the vendor-specific complexity. Phase 3 stages '
    'are priced individually as we reach them. We don\u2019t want to guess at numbers \u2014 we want '
    'to give you ones we can stand behind.'
)

doc.add_heading('Ongoing support', level=3)
add_paragraph(
    'This is the part that\u2019s easy to overlook and important to get right.'
)
add_paragraph(
    'Your team does exceptional work running the business. But the reality is that a custom software '
    'solution \u2014 no matter how well built \u2014 requires ongoing technical support. Security '
    'patches, vendor format changes, bug fixes, feature requests, and day-to-day questions from '
    'staff who are learning a new process.'
)
add_paragraph(
    'The support commitment will grow as the system does. After Phase 2 launches, we anticipate '
    '[X\u2013X hours/month] for the card generation tool. Once the order platform is live, that '
    'increases to [X\u2013X hours/month] to cover a larger surface area \u2014 customer-facing '
    'systems, QuickBooks integration, security monitoring, and day-to-day staff support. The first '
    'quarter after each launch will require heavier involvement as your team builds comfort.'
)
add_paragraph(
    'This is a restructuring of our existing support arrangement, not an entirely new cost \u2014 '
    'but it will be more than what the current retainer covers, and we want that to be clear upfront.'
)
add_paragraph(
    '**The alternative \u2014 building something and walking away \u2014 doesn\u2019t work for a '
    'team at your current technical capacity.** Every hour we spend supporting your operations is '
    'an hour your team isn\u2019t stuck, and that\u2019s the point.'
)

doc.add_heading('The growth case', level=3)
add_paragraph(
    '[Revenue/volume numbers to be confirmed]'
)
add_paragraph(
    'Your digital card business currently represents approximately [X%] of total revenue. If digital '
    'order volume doubles over the next [X] years \u2014 which the industry trend suggests is '
    'conservative \u2014 the current process can\u2019t handle it. The investment in Phases 1\u20133 '
    'isn\u2019t just about fixing what\u2019s broken today. It\u2019s about building the infrastructure '
    'that lets the digital side of your business grow without hitting a wall.'
)

add_horizontal_rule()

# --- NEXT STEPS ---
doc.add_heading('Next Steps', level=2)
add_paragraph('Two things need to happen to get this moving.')
p = doc.add_paragraph(style='List Number')
add_bold_text(p, 'Walk through this document together.')
p.add_run(
    ' We\u2019ll schedule a call to align on priorities, answer questions, and get a quick update '
    'on the Save-on-Foods and Sequoia conversations \u2014 not because it changes Phase 1, but '
    'because it helps us architect for what\u2019s ahead.'
)
p = doc.add_paragraph(style='List Number')
add_bold_text(p, 'Schedule the Lloyd hand-off session.')
p.add_run(
    ' Redstamp sends an engineer to your office in April. This is the single action that unblocks '
    'everything else \u2014 and Phase 1 is the smallest commitment with the highest information value.'
)

# --- Save ---
output_path = r'D:\projects-work\rs-progressive-fundraising\projects\gift-cards\docs\plans\Progressive Gift Cards — Where We Go From Here.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
