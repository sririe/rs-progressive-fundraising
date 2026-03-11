"""
Generate a .docx for the internal solution comparison document.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# --- Style configuration ---
style = doc.styles['Normal']
font = style.font
font.name = 'Inter'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Inter'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
    hs.paragraph_format.space_after = Pt(6)
    hs.font.size = Pt([0, 24, 18, 14][level])

# --- Helpers ---

def add_bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run

def add_paragraph(text):
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def add_table(rows, header=True):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            parts = re.split(r'(\*\*.*?\*\*)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = 'Inter'
            if header and i == 0:
                for run in p.runs:
                    run.bold = True
                shading = cell._element.get_or_add_tcPr()
                elem = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'F2F2F2', qn('w:val'): 'clear'})
                shading.append(elem)
    doc.add_paragraph()
    return table

def add_kv_table(pairs):
    table = doc.add_table(rows=len(pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    for i, (key, value) in enumerate(pairs):
        ck = table.rows[i].cells[0]
        ck.text = ''
        p = ck.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(key)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Inter'
        shading = ck._element.get_or_add_tcPr()
        elem = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'F8F8F8', qn('w:val'): 'clear'})
        shading.append(elem)
        cv = table.rows[i].cells[1]
        cv.text = ''
        p = cv.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.name = 'Inter'
    doc.add_paragraph()
    return table

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): 'CCCCCC'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_option(name, what_it_is, tech_stack, team_comfort, numbers, pros, cons, capacity):
    doc.add_heading(name, level=3)
    add_paragraph(f'**What it is:** {what_it_is}')
    add_paragraph(f'**Tech stack:** {tech_stack}')
    add_paragraph(f'**Team comfort:** {team_comfort}')
    add_kv_table(numbers)
    add_paragraph('**Pros:**')
    for pro in pros:
        add_bullet(pro)
    add_paragraph('**Cons:**')
    for con in cons:
        add_bullet(con)
    add_paragraph(f'**Redstamp capacity impact:** {capacity}')
    add_hr()


# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Title
title = doc.add_heading('Card Generation Tool \u2014 Internal Solution Comparison', level=1)
p = doc.add_paragraph()
run = p.add_run('For team discussion. Not client-facing.')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
run = p.add_run('Redstamp  |  March 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p.paragraph_format.space_after = Pt(16)

add_paragraph(
    'This covers the card generation / fulfillment tool only \u2014 not the '
    'customer portal, white-label, or order matching work.'
)

# --- The Problem ---
doc.add_heading('The Problem We\u2019re Solving', level=2)
add_paragraph(
    'Progressive\u2019s digital card generation is a manual pipeline run by one person (Lloyd) '
    'using custom scripts on a local machine. The process varies by vendor, involves sensitive '
    'card data, and can\u2019t be picked up by non-technical staff. Lloyd wants out. Mario is '
    'the designated successor. The solution needs to be simple enough that Mario \u2014 or a '
    'new hire \u2014 can run it without developer support.'
)

# --- Three Approaches ---
doc.add_heading('Three Approaches', level=2)

add_option(
    'Option 1: Guided Desktop Tool (Electron)',
    'A local application on their office machine. Select vendor, drop in spreadsheet, '
    'generate card URLs. Card data never leaves their network.',
    'JavaScript, HTML/CSS, Node.js. Electron for desktop packaging. SheetJS for spreadsheet handling.',
    'High. Web technologies the dev team already works in, deployed differently.',
    [
        ('Build estimate (Redstamp)', '$8K\u2013$15K'),
        ('Build hours', '80\u2013150 hrs'),
        ('Market equivalent', '$25K\u2013$55K'),
        ('Monthly ongoing (stabilization, Q1)', '8\u201312 hrs/mo'),
        ('Monthly ongoing (steady state)', '3\u20135 hrs/mo'),
        ('Monthly cost to Progressive (steady state)', '~$500\u2013$750/mo'),
    ],
    [
        'Simplest security model \u2014 no web exposure',
        'Fastest to V1',
        'Lowest ongoing cost',
        'Team can build and maintain confidently',
    ],
    [
        'Harder to support remotely when something breaks',
        'Tied to a physical machine',
        'No visibility for Redstamp into day-to-day operations',
        'Limited extensibility \u2014 connecting this to a portal or adding features later means rearchitecting',
    ],
    'Low. Minimal ongoing drain once stable.',
)

add_option(
    'Option 2: Secure Web Application (Laravel)',
    'A browser-based tool hosted in a secure environment. Same workflow as Option 1, but '
    'accessible from anywhere. Redstamp can monitor, update, and support it remotely.',
    'Laravel (PHP), PostgreSQL, hosted on Forge/DigitalOcean or Railway. Built-in auth and '
    'encryption. Blade templates for the frontend.',
    'Moderate. PHP is familiar from WordPress. Laravel is a step up in structure but the '
    'language is the same. The hosting and deployment patterns are new.',
    [
        ('Build estimate (Redstamp)', '$18K\u2013$35K'),
        ('Build hours', '185\u2013305 hrs'),
        ('Market equivalent', '$55K\u2013$120K'),
        ('Monthly ongoing (stabilization, Q1)', '12\u201318 hrs/mo'),
        ('Monthly ongoing (steady state)', '5\u201310 hrs/mo'),
        ('Monthly cost to Progressive (steady state)', '~$750\u2013$1,500/mo'),
    ],
    [
        'Easiest to maintain and update long-term',
        'Redstamp has full remote access for support',
        'Natural path to portal integration, new vendors, additional features',
        'Laravel\u2019s security patterns are mature and well-documented',
        'Extensive documentation and community patterns for tooling support',
    ],
    [
        'Card data transits the internet \u2014 security architecture must be solid',
        'More upfront investment',
        'Hosting costs (modest \u2014 $20\u201350/mo for infrastructure)',
        'Tim\u2019s \u201cbank site\u201d concern applies \u2014 needs to be addressed head-on in the security design',
    ],
    'Moderate. Ongoing hosting management, security patches, and support. But remote access makes it efficient \u2014 no site visits needed.',
)

add_option(
    'Option 3: Hybrid \u2014 Web Workflow + Isolated Processing',
    'Web interface for order management and workflow. Actual card data processing happens in a '
    'sandboxed environment (serverless functions or containers) that spins up, generates URLs, '
    'and tears down. The web app never touches raw card numbers.',
    'Laravel or Next.js for web layer. AWS Lambda or Cloudflare Workers for processing. API '
    'layer connecting the two. Encrypted temporary storage for card data in transit.',
    'Low to moderate. The web layer is familiar enough (especially if Laravel). The '
    'serverless/container infrastructure is new territory for the team. Debugging failures in '
    'the processing layer requires concepts they haven\u2019t worked with.',
    [
        ('Build estimate (Redstamp)', '$28K\u2013$50K'),
        ('Build hours', '260\u2013425 hrs'),
        ('Market equivalent', '$80K\u2013$170K'),
        ('Monthly ongoing (stabilization, Q1)', '15\u201322 hrs/mo'),
        ('Monthly ongoing (steady state)', '8\u201314 hrs/mo'),
        ('Monthly cost to Progressive (steady state)', '~$1,200\u2013$2,000/mo'),
    ],
    [
        'Strongest security boundary \u2014 cleanest separation of sensitive data from the web interface',
        'Most extensible architecture for future growth',
        'Easiest to pass a formal security review if Progressive ever needs one',
    ],
    [
        'Most complex to build and architect',
        'Longest time to V1',
        'Highest ongoing maintenance burden',
        'Team would struggle to debug the infrastructure layer independently',
        'Overkill for current volume \u2014 this architecture is designed for scale Progressive may not reach for years',
    ],
    'High. Redstamp is effectively the ops team for the infrastructure layer indefinitely. This is a real ongoing commitment.',
)

# --- Recommendation ---
doc.add_heading('Recommendation for What to Pitch', level=2)

add_table([
    ['Factor', 'Option 1 (Desktop)', 'Option 2 (Web App)', 'Option 3 (Hybrid)'],
    ['Solves the Lloyd problem', 'Yes', 'Yes', 'Yes'],
    ['Mario can run it', 'Yes', 'Yes', 'Yes, but debugging is harder'],
    ['Team can build it', 'Confidently', 'With some stretch', 'Significant stretch'],
    ['Team can maintain it', 'Yes', 'Yes, with learning curve', 'Needs Redstamp ongoing'],
    ['Extends to portal later', 'Poorly', 'Naturally', 'Naturally'],
    ['Monthly Redstamp commitment', '3\u20135 hrs', '5\u201310 hrs', '8\u201314 hrs'],
    ['Client monthly cost', '~$500\u2013750', '~$750\u20131,500', '~$1,200\u20132,000'],
])

add_paragraph(
    '**Option 3 is probably too much.** The security benefit is real but the ongoing capacity '
    'commitment is high, the team can\u2019t own it independently, and the architecture solves '
    'for scale Progressive doesn\u2019t need yet. It\u2019s the right answer for a company with '
    'a dedicated engineering team. Progressive isn\u2019t that company.'
)
add_paragraph(
    '**Option 1 is the fastest win but the worst foundation.** It solves the immediate Lloyd '
    'problem but creates a new island \u2014 a local app that\u2019s hard to extend, hard to '
    'support remotely, and disconnected from everything else we\u2019d build. If the plan is to '
    'eventually connect this to a customer portal, Option 1 becomes throwaway work.'
)
add_paragraph(
    '**Option 2 is the sweet spot.** It solves the immediate problem, the team can build it in '
    'a stack adjacent to what they know, Redstamp can support it remotely, and it\u2019s the '
    'natural foundation for the portal and any future features. The ongoing commitment is '
    'manageable. The security concern is real but solvable with the right design \u2014 and the '
    'Lloyd session will inform exactly how card data needs to be handled.'
)

doc.add_heading('What to put in front of Doug', level=3)
add_paragraph(
    'Propose Option 2 as the recommended path. Reference the market cost ($55K\u2013$120K from '
    'a traditional firm) to anchor the value. Price the build at [TBD \u2014 needs internal '
    'pricing session]. Include the ongoing monthly support as a line item so there are no surprises.'
)
add_paragraph(
    'Option 1 could be mentioned as a faster, lower-cost alternative if Doug pushes back on '
    'investment \u2014 but position it honestly as a shorter-term fix that doesn\u2019t set up '
    'the next phase.'
)
add_paragraph('Option 3 doesn\u2019t need to be in the proposal. It\u2019s overbuilt for where they are.')

add_hr()

# --- Order Platform ---
doc.add_heading('Order Platform Re-architecture (Phase 3 in the Proposal)', level=2)
add_paragraph(
    'Separate from the card generation tool. This replaces the Formidable Forms setup with a '
    'proper B2B order platform: customer accounts, 2FA, order submission, order history/status, '
    'admin backend, QuickBooks integration, Benji Pays workflow, white-label ready.'
)

add_table([
    ['', 'Traditional', 'Redstamp'],
    ['**Total hours**', '610\u20131,050', '285\u2013490'],
    ['**Market rate ($150\u2013200/hr)**', '$90K\u2013$210K', '\u2014'],
    ['**Redstamp**', '\u2014', '$30K\u2013$60K'],
])

add_paragraph('**Phased internally:**')
add_table([
    ['Stage', 'Scope', 'Redstamp estimate'],
    ['MVP', 'Accounts, 2FA, order submission, status, admin view, QuickBooks integration', '$18K\u2013$30K'],
    ['Digital delivery', 'Card retrieval connected to gen tool, direct-to-recipient', '$8K\u2013$15K'],
    ['White-label', 'Branded vendor portals, white-label architecture', '$8K\u2013$18K'],
])

add_paragraph(
    '**Ongoing support (steady state):** 8\u201315 hrs/month (~$1,200\u2013$2,200/mo) \u2014 '
    'higher than the card generation tool because it\u2019s customer-facing and touches payment flows.'
)
add_paragraph(
    '**Stack:** Same as Phase 2 (likely Next.js). Shared codebase or shared infrastructure would '
    'reduce the combined maintenance burden. An outside consultant for initial architecture and '
    'security is worth considering here as well.'
)

add_hr()

# --- Stack Note ---
doc.add_heading('Stack Note', level=2)
add_paragraph(
    'The team has some React experience, and Spencer has successfully shipped React projects '
    'using our current development workflows. This means a **Next.js** path for Option 2 may be '
    'more practical than Laravel despite Laravel being closer to the WordPress/PHP world. Next.js '
    'also has the strongest tooling ecosystem (Vercel\u2019s own tools, extensive documentation, '
    'large community pattern library).'
)
add_paragraph(
    'If we go Next.js, the main gap would be backend patterns and security architecture \u2014 '
    'not the React frontend. An **outside consultant** for the initial architecture and security '
    'setup could bridge that gap, with the team owning ongoing development and maintenance once '
    'the foundation is solid.'
)
add_paragraph(
    'This changes the team comfort assessment for Option 2 from \u201cmoderate\u201d to '
    '\u201cmoderate-high\u201d if we go the Next.js route.'
)

# --- Open Questions ---
doc.add_heading('Open Questions for Team Discussion', level=2)
questions = [
    'Confirm React comfort level with Tim and Bronte \u2014 is it enough to maintain a Next.js app, or limited to component-level work?',
    'Does the dev team have any Laravel experience, or is this a full ramp-up?',
    'What\u2019s the realistic timeline for the Lloyd session? Can we schedule it within the next 2 weeks?',
    'Do we have access to Progressive\u2019s QuickBooks to pull revenue/volume numbers for the proposal?',
    'What\u2019s our current retainer structure with Progressive, and what would a restructured support agreement look like?',
    'Has Doug mentioned any budget ceiling or range in recent conversations?',
]
for q in questions:
    add_bullet(q)

# --- Save ---
output_path = r'D:\projects-work\rs-progressive-fundraising\docs\plans\Internal Solution Comparison — Card Generation Tool.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
